#!/usr/bin/env python3
"""Colore em vermelho e negrito os marcadores de revisão de uma peça .docx.

As skills de redação da Letra da Lei (peticao, contestacao, fundamentacao,
analise) inserem marcadores inline que o(a) advogado(a)/magistrado(a) precisa
resolver antes de protocolar/publicar. Este script percorre o documento e
aplica cor #FF0000 + negrito a cada ocorrência desses marcadores, para que
fiquem visualmente óbvios na revisão.

Uso:
    python scripts/colorir_marcadores.py caminho/para/peca.docx

O arquivo é sobrescrito no mesmo caminho.

Dependência: python-docx  (pip install python-docx)

Marcadores reconhecidos: qualquer trecho entre colchetes cujo conteúdo comece
por uma das palavras-chave conhecidas (VERIFICAR, CITAÇÃO PENDENTE, FORA DO
CORPUS, JURISPRUDÊNCIA, DOC. A NUMERAR, CONVICÇÃO JUDICIAL, A REVISAR, TEMA
PENDENTE, ATENÇÃO). Colchetes comuns de preenchimento (ex.: "[Cidade]",
"[NOME DO AUTOR]") NÃO são tocados.

Limitação conhecida: um marcador precisa estar contido em um único "run" do
.docx (o caso normal quando o texto é gerado pela skill). Marcadores quebrados
entre runs com formatações diferentes podem não ser coloridos; nesse caso,
aplique a cor manualmente.
"""

import copy
import re
import sys

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor
    from docx.text.run import Run
except ImportError:
    sys.exit(
        "Erro: python-docx não está instalado. Rode: pip install python-docx\n"
        "Ou aplique a cor #FF0000 manualmente aos marcadores entre colchetes."
    )

RED = RGBColor(0xFF, 0x00, 0x00)

# Palavras-chave que identificam um marcador (sem acento e com acento cobertos).
KEYWORDS = [
    r"VERIFICAR",
    r"CITA[ÇC][ÃA]O PENDENTE",
    r"FORA DO CORPUS",
    r"JURISPRUD[ÊE]NCIA",
    r"DOC\.? A NUMERAR",
    r"CONVIC[ÇC][ÃA]O JUDICIAL",
    r"A REVISAR",
    r"TEMA PENDENTE",
    r"ATEN[ÇC][ÃA]O",
]
KEYWORD_RE = re.compile("|".join(KEYWORDS), re.IGNORECASE)
# Um marcador é um colchete cujo conteúdo (após espaços iniciais) bate uma keyword.
BRACKET_RE = re.compile(r"\[[^\[\]]*\]")


def _is_marker(bracket_text: str) -> bool:
    return bool(KEYWORD_RE.search(bracket_text))


def _clone_run(src_r, text: str):
    """Clona um <w:r> preservando rPr, substituindo o texto."""
    new_r = copy.deepcopy(src_r)
    for child in list(new_r):
        if child.tag in (qn("w:t"), qn("w:br"), qn("w:tab"), qn("w:cr")):
            new_r.remove(child)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_r.append(t)
    return new_r


def _color_paragraph(paragraph) -> int:
    """Colore marcadores nos runs de um parágrafo. Retorna nº de marcadores."""
    count = 0
    for run in list(paragraph.runs):
        text = run.text
        if not text or not KEYWORD_RE.search(text):
            continue
        spans = [(m.start(), m.end()) for m in BRACKET_RE.finditer(text) if _is_marker(m.group(0))]
        if not spans:
            continue

        # Quebra o texto do run em segmentos (texto normal vs. marcador).
        segments = []
        idx = 0
        for start, end in spans:
            if start > idx:
                segments.append((text[idx:start], False))
            segments.append((text[start:end], True))
            idx = end
        if idx < len(text):
            segments.append((text[idx:], False))

        src_r = run._r
        parent = src_r.getparent()
        for seg_text, is_marker in segments:
            new_r = _clone_run(src_r, seg_text)
            src_r.addprevious(new_r)
            if is_marker:
                wrapped = Run(new_r, paragraph)
                wrapped.font.color.rgb = RED
                wrapped.font.bold = True
                count += 1
        parent.remove(src_r)
    return count


def _iter_paragraphs(container):
    """Percorre parágrafos do corpo e de tabelas (recursivo)."""
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def main(argv):
    if len(argv) != 2:
        sys.exit("Uso: python scripts/colorir_marcadores.py caminho/para/peca.docx")
    path = argv[1]
    try:
        doc = Document(path)
    except Exception as exc:  # noqa: BLE001
        sys.exit(f"Não foi possível abrir '{path}': {exc}")

    total = sum(_color_paragraph(p) for p in _iter_paragraphs(doc))
    doc.save(path)
    print(f"{total} marcador(es) colorido(s) em vermelho/negrito → {path}")
    if total == 0:
        print("Nenhum marcador encontrado. Confira se a peça realmente os contém.")


if __name__ == "__main__":
    main(sys.argv)
