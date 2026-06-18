#!/usr/bin/env python3
"""Utilitários para montar peças .docx no padrão forense da Letra da Lei.

Centraliza a formatação compartilhada pelas skills de redação (peticao,
contestacao, fundamentacao, analise), para que cada peça NÃO reimplemente o
mesmo código python-docx. Padrão: corpo justificado (Times New Roman 12pt),
blocos de citação recuados 1,25 cm em 11pt (sem aspas), com a linha "Fonte:"
logo abaixo e uma linha em branco depois.

Uso (no seu script de geração da peça):

    import sys; sys.path.insert(0, "${CLAUDE_PLUGIN_ROOT}/scripts")
    from montar_docx import nova_peca, paragrafo, titulo, bloco_citacao, salvar

    doc = nova_peca()
    titulo(doc, "AÇÃO DE INDENIZAÇÃO POR DANOS MORAIS")
    paragrafo(doc, "I, DOS FATOS", negrito=True)
    paragrafo(doc, "No dia [VERIFICAR: data], a parte autora ...")
    bloco_citacao(
        doc,
        "Art. 14. O fornecedor de serviços responde, independentemente da "
        "existência de culpa, pela reparação dos danos causados ...",
        "Fonte: CDC, art. 14 | https://www.planalto.gov.br/... | situação: vigente | lei: Lei-8078-1990",
    )
    salvar(doc, "outputs/peticao-danos-morais-voo-2026-06-18.docx")

Depois de salvar, rode colorir_marcadores.py no arquivo para destacar em
vermelho/negrito os marcadores de revisão.

Dependência: python-docx (pip install python-docx).
"""
import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

CORPO_PT = 12      # corpo do texto
RECUO_CM = 1.25    # recuo do bloco de citação (≈720 DXA)
FONTE = "Times New Roman"


def nova_peca(corpo_pt: int = CORPO_PT, fonte: str = FONTE) -> Document:
    """Documento novo com o estilo Normal já justificado, na fonte do padrão."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = fonte
    normal.font.size = Pt(corpo_pt)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return doc


def paragrafo(doc: Document, texto: str, negrito: bool = False, centralizado: bool = False):
    """Parágrafo de corpo (justificado por padrão)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centralizado else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texto)
    run.bold = negrito
    return p


def titulo(doc: Document, texto: str):
    """Título centralizado em negrito (ex.: a denominação da ação)."""
    return paragrafo(doc, texto, negrito=True, centralizado=True)


def bloco_citacao(doc: Document, texto: str, fonte_linha: str, corpo_pt: int = CORPO_PT):
    """Transcrição recuada 1,25 cm, 1pt menor, sem aspas; + linha 'Fonte:'; + linha em branco.

    `fonte_linha` deve começar com 'Fonte:' e conter citacao + source_url (lei)
    ou autoridade/tipo/search_id/eficacia (jurisprudência), conforme
    _shared/citacao-e-formato.md.
    """
    menor = Pt(corpo_pt - 1)
    for linha in (texto, fonte_linha):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(RECUO_CM)
        run = p.add_run(linha)
        run.font.size = menor
    doc.add_paragraph()  # linha em branco separando do texto seguinte


def salvar(doc: Document, caminho: str) -> str:
    """Cria o diretório se preciso e salva. Retorna o caminho."""
    os.makedirs(os.path.dirname(caminho) or ".", exist_ok=True)
    doc.save(caminho)
    return caminho
